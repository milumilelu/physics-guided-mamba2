"""Verify V6 template format consistency."""
from docx import Document

doc = Document(r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V6_模板格式版.docx")

print("=" * 70)
print("V6 TEMPLATE FORMAT VERIFICATION")
print("=" * 70)

# 1. Document title
print("\n--- 1. Document Title ---")
first_para = doc.paragraphs[0]
print(f"  Text: {first_para.text}")
print(f"  Alignment: {first_para.alignment}")
print(f"  Bold: {first_para.runs[0].bold if first_para.runs else 'N/A'}")

# 2. Header info table
print("\n--- 2. Header Info Table ---")
if doc.tables:
    t = doc.tables[0]
    print(f"  Rows: {len(t.rows)}, Cols: {len(t.columns)}")
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip()[:35] for c in row.cells]
        print(f"  Row {ri}: {' | '.join(cells)}")

# 3. Section headings (template style)
print("\n--- 3. Section Headings (template question style) ---")
expected_headings = [
    "术语解释",
    "1、本发明要解决的技术问题是什么？",
    "2、详细介绍技术背景",
    "3、以因果关系推理的方式推导出现有技术的缺点",
    "4、本发明技术方案的详细阐述",
    "5、本发明的关键点和欲保护点是什么？",
    "6、能实现本发明目的的其他替代方案",
    "参考文献",
    "附录",
]
headings_found = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and any(text.startswith(h[:6]) for h in expected_headings):
        if p.runs and p.runs[0].bold:
            headings_found.append((i, text[:80]))
            print(f"  [{i:3d}] {text[:80]}")

# 4. Section 4 internal structure
print("\n--- 4. Section 4 Internal Structure ---")
in_sec4 = False
sec4_items = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "4、本发明技术方案的详细阐述" in text:
        in_sec4 = True
        continue
    if in_sec4:
        if text.startswith("5、"):
            break
        if text.startswith("4.1") or text.startswith("4.2") or text.startswith("4.3") or text.startswith("4.4"):
            sec4_items.append(("NOTE", text[:60]))
        elif text.startswith("（1）") or text.startswith("（2）") or text.startswith("（3）") or text.startswith("（4）") or text.startswith("（5）") or text.startswith("（6）") or text.startswith("（7）"):
            sec4_items.append(("ITEM", text[:60]))
        elif "结合附图" in text:
            sec4_items.append(("INTRO", text[:60]))

for item_type, text in sec4_items:
    print(f"  [{item_type:5s}] {text}")

# 5. References format
print("\n--- 5. References Format ---")
in_refs = False
ref_count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith("参考文献"):
        in_refs = True
        continue
    if in_refs and text:
        if text.startswith("附录"):
            break
        if text.startswith("【"):
            ref_count += 1
            print(f"  {text[:70]}")
print(f"  Total references: {ref_count}")

# 6. Content statistics
print("\n--- 6. Content Statistics ---")
total_paras = len(doc.paragraphs)
total_tables = len(doc.tables)
total_chars = sum(len(p.text) for p in doc.paragraphs)
non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
print(f"  Total paragraphs: {total_paras}")
print(f"  Non-empty paragraphs: {non_empty}")
print(f"  Total tables: {total_tables}")
print(f"  Total characters: {total_chars}")

# 7. Check for ablation content (should be empty)
print("\n--- 7. Ablation Content Check (should be empty) ---")
ablation_found = False
for i, p in enumerate(doc.paragraphs):
    text = p.text
    for kw in ["消融", "M0", "M1", "M2", "M3", "M4", "4.9"]:
        if kw in text:
            if kw in ["M1", "M2", "M3", "M4"] and "模型" not in text and "消融" not in text:
                continue
            print(f"  WARNING [{i}] keyword='{kw}': {text[:60]}")
            ablation_found = True
if not ablation_found:
    print("  PASS: No ablation content found.")

# 8. Check font
print("\n--- 8. Font Check ---")
sample_paras = [p for p in doc.paragraphs if p.text.strip() and p.runs][:5]
for i, p in enumerate(sample_paras):
    run = p.runs[0]
    print(f"  Para {i}: font={run.font.name}, size={run.font.size}, bold={run.bold}, text={p.text[:30]}")

print("\n" + "=" * 70)
print("VERIFICATION COMPLETE")
print("=" * 70)
