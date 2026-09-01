"""
Convert V5 optimized patent disclosure to standard template format (V6).
Template reference: 专利技术交底书1.doc (照相拼图工件定位)
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.parts.image import ImagePart
import copy

V5_PATH = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V5_优化版.docx"
V6_PATH = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V6_模板格式修复版.docx"

# Read V5 content
v5 = Document(V5_PATH)
v5_paras = v5.paragraphs
v5_tables = v5.tables

# Extract content by section from V5
def extract_section(paras, start_heading, end_heading=None):
    """Extract paragraphs between two headings."""
    result = []
    in_section = False
    for p in paras:
        if p.style.name == "Heading 1" and start_heading in p.text:
            in_section = True
            continue
        if in_section:
            if end_heading and p.style.name == "Heading 1" and end_heading in p.text:
                break
            if p.style.name == "Heading 1" and p.text != start_heading:
                break
            result.append(p)
    return result

def extract_subsection(paras, heading_text):
    """Extract paragraphs under a Heading 2 subsection."""
    result = []
    in_section = False
    for p in paras:
        if p.style.name == "Heading 2" and heading_text in p.text:
            in_section = True
            continue
        if in_section:
            if p.style.name in ("Heading 1", "Heading 2"):
                break
            result.append(p)
    return result

# Extract V5 sections
sec1 = extract_section(v5_paras, "1、本发明要解决的技术问题")
sec2 = extract_section(v5_paras, "2、技术背景及最接近现有技术")
sec3 = extract_section(v5_paras, "3、现有技术的缺点")
sec4 = extract_section(v5_paras, "4、本发明技术方案的详细阐述")
sec5 = extract_section(v5_paras, "5、本发明的关键点和欲保护点")
sec6 = extract_section(v5_paras, "6、能实现本发明目的的其他替代方案")
appendix = extract_section(v5_paras, "附录A：代理人参考材料")
references = extract_section(v5_paras, "参考资料")

# Extract 4.x subsections
sub41 = extract_subsection(v5_paras, "4.1 总体技术路线与术语定义")
sub42 = extract_subsection(v5_paras, "4.2 输入变量及双状态定义")
sub43 = extract_subsection(v5_paras, "4.3 历史状态驱动的材料去除状态门控")
sub44 = extract_subsection(v5_paras, "4.4 压缩观测与多尺度联合约束")
sub45 = extract_subsection(v5_paras, "4.5 训练目标与算法实现")
sub46 = extract_subsection(v5_paras, "4.6 现有实验数据与可实施性")
sub47 = extract_subsection(v5_paras, "4.7 预测输出与工程应用")

# Extract term definitions from 4.1
term_defs = []
for p in sub41:
    if "——" in p.text and any(kw in p.text for kw in ["表面形貌状态", "脉冲历史材料状态", "材料去除状态门控", "压缩观测", "多尺度联合约束"]):
        term_defs.append(p.text)

# ============================================================
# Create V6 document with template format
# ============================================================
doc = Document()

# Set default font to 宋体
style = doc.styles["Normal"]
font = style.font
font.name = "宋体"
font.size = Pt(12)  # 小四号
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# Set page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_title(text, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12):
    """Add a title paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_heading_text(text, size=14, bold=True, space_before=12, space_after=6):
    """Add a section heading (template style - no Heading style, just bold text)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_body(text, indent=True, bold=False, size=12):
    """Add a body paragraph with first-line indent."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_list_item(text, size=12):
    """Add a list item with (1) (2) numbering, first line indent."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def copy_paragraph_content(src_para, indent=True):
    """Copy paragraph content including text, formatting, and images via XML deep copy."""
    has_content = bool(src_para.text.strip()) or bool(src_para._element.findall('.//' + qn('w:drawing')))
    if not has_content:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        return p
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    # Deep copy all child elements except pPr (paragraph properties)
    for child in src_para._element:
        if child.tag == qn('w:pPr'):
            continue
        new_child = copy.deepcopy(child)
        p._element.append(new_child)
    # Fix image relationships: find all blip elements and update r:embed
    for blip in p._element.findall('.//' + qn('a:blip')):
        old_rId = blip.get(qn('r:embed'))
        if old_rId and old_rId in src_para.part.rels:
            src_rel = src_para.part.rels[old_rId]
            src_image_part = src_rel.target_part
            # Create new image part in destination document
            new_image_part = ImagePart(
                src_image_part.partname,
                src_image_part.content_type,
                src_image_part.blob,
                doc.part.package
            )
            # Add relationship and get new rId
            new_rId = doc.part.relate_to(new_image_part, src_rel.reltype)
            blip.set(qn('r:embed'), new_rId)
    # Ensure font is set on all runs
    for run in p.runs:
        if run.font.name is None:
            run.font.name = "宋体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if run.font.size is None:
            run.font.size = Pt(12)
    return p

def copy_paragraphs(para_list, indent=True):
    """Copy a list of paragraphs."""
    for p in para_list:
        copy_paragraph_content(p, indent=indent)

# ============================================================
# 1. Document title
# ============================================================
add_title("专 利 申 请 技 术 交 底 书", size=22, space_after=16)

# ============================================================
# 2. Header info table (template format)
# ============================================================
info_table = doc.add_table(rows=10, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = "Table Grid"

# Set column widths
for row in info_table.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(5)
    row.cells[2].width = Cm(3)
    row.cells[3].width = Cm(5)

def set_cell(cell, text, bold=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# Row 0: 发明类型
set_cell(info_table.rows[0].cells[0], "发明类型：", bold=True)
set_cell(info_table.rows[0].cells[1], "发明 ☑  实用新型 ☐  外观设计 ☐")
info_table.rows[0].cells[1].merge(info_table.rows[0].cells[3])

# Row 1: 同日申请 + 请求提前公布
set_cell(info_table.rows[1].cells[0], "同日申请实用新型：", bold=True)
set_cell(info_table.rows[1].cells[1], "☐")
set_cell(info_table.rows[1].cells[2], "请求提前公布：", bold=True)
set_cell(info_table.rows[1].cells[3], "☐")

# Row 2: 优先权/实审方式
set_cell(info_table.rows[2].cells[0], "优先权 实审方式：", bold=True)
set_cell(info_table.rows[2].cells[1], "☐ 新申请提  ☐ 公布提  ☐ 期限届满提  ☐ 不提")
info_table.rows[2].cells[1].merge(info_table.rows[2].cells[3])

# Row 3: 发明名称
set_cell(info_table.rows[3].cells[0], "发明名称：", bold=True)
set_cell(info_table.rows[3].cells[1], "一种基于双状态机制门控及多尺度观测约束的超快激光多脉冲烧蚀形貌预测方法")
info_table.rows[3].cells[1].merge(info_table.rows[3].cells[3])

# Row 4: 发明人 + 身份证号
set_cell(info_table.rows[4].cells[0], "发明人：", bold=True)
set_cell(info_table.rows[4].cells[1], "（待填写）")
set_cell(info_table.rows[4].cells[2], "身份证号：", bold=True)
set_cell(info_table.rows[4].cells[3], "（待填写）")

# Row 5: 技术联系人 + 所属公司/部门
set_cell(info_table.rows[5].cells[0], "技术联系人：", bold=True)
set_cell(info_table.rows[5].cells[1], "（待填写）")
set_cell(info_table.rows[5].cells[2], "所属公司/部门：", bold=True)
set_cell(info_table.rows[5].cells[3], "（待填写）")

# Row 6: 电话 + 传真
set_cell(info_table.rows[6].cells[0], "电话：", bold=True)
set_cell(info_table.rows[6].cells[1], "（待填写）")
set_cell(info_table.rows[6].cells[2], "传真：", bold=True)
set_cell(info_table.rows[6].cells[3], "（待填写）")

# Row 7: E-mail
set_cell(info_table.rows[7].cells[0], "E-mail：", bold=True)
set_cell(info_table.rows[7].cells[1], "（待填写）")
info_table.rows[7].cells[1].merge(info_table.rows[7].cells[3])

# Row 8: 版本说明
set_cell(info_table.rows[8].cells[0], "版本说明：", bold=True)
set_cell(info_table.rows[8].cells[1], "V6：标准模板格式版（对齐专利交底书模板格式，删除消融实验内容）")
info_table.rows[8].cells[1].merge(info_table.rows[8].cells[3])

# Row 9: empty spacer
for cell in info_table.rows[9].cells:
    set_cell(cell, "")

doc.add_paragraph()  # spacer

# ============================================================
# 3. 术语解释
# ============================================================
add_heading_text("术语解释：", size=14, space_before=6)
for term in term_defs:
    add_body(term, indent=True)

# ============================================================
# 4. Section 1 - 技术问题
# ============================================================
add_heading_text("1、本发明要解决的技术问题是什么？", size=14)
copy_paragraphs(sec1, indent=True)

# ============================================================
# 5. Section 2 - 技术背景
# ============================================================
add_heading_text("2、详细介绍技术背景,并描述已有的与本发明最相近似的实现方案（包括两部分：1、作为本发明基础的且帮助理解本发明公知技术内容；2、与本发明最接近的已有技术方案的说明（对于方法，应说明现有方法的步骤，对于装置，应当说明结构组成及其关系））", size=14)

# 2.1 content
add_heading_text("2.1 从工程现象到物理建模基础", size=12, space_before=8)
sub21 = extract_subsection(v5_paras, "2.1 从工程现象到物理建模基础")
copy_paragraphs(sub21, indent=True)

# 2.2 content
add_heading_text("2.2 与本发明最接近的公开方案及查新判断", size=12, space_before=8)
sub22 = extract_subsection(v5_paras, "2.2 与本发明最接近的公开方案及查新判断")
copy_paragraphs(sub22, indent=True)

# Copy 查新对比表 (Table 1 from V5)
if len(v5_tables) >= 2:
    doc.add_paragraph()
    src_table = v5_tables[1]  # 查新对比表
    new_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    new_table.style = "Table Grid"
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell(new_table.rows[ri].cells[ci], cell.text, size=10)
    doc.add_paragraph()

# 2.3 content
add_heading_text("2.3 中国人工智能相关专利撰写考虑", size=12, space_before=8)
sub23 = extract_subsection(v5_paras, "2.3 中国人工智能相关专利撰写考虑")
copy_paragraphs(sub23, indent=True)

# ============================================================
# 6. Section 3 - 缺点/目的/效果
# ============================================================
add_heading_text("3、以因果关系推理的方式推导出现有技术的缺点是什么？针对这些缺点，说明本发明的目的及能够达到的技术效果。（现有技术的缺点是针对于本发明的优点来说的，本发明不能解决的缺点不必写）", size=14)
copy_paragraphs(sec3, indent=True)

# ============================================================
# 7. Section 4 - 技术方案详细阐述 (template format)
# ============================================================
add_heading_text("4、本发明技术方案的详细阐述，应该结合流程图、原理框图、电路图、系统结构图进行说明（发明中每一功能的实现都要有相应的技术实现方案；所有附图都应该有详细的文字描述；方法专利都应该提供流程图，并提供相关的系统装置）。", size=14)

# Template filling requirements (no numbering to avoid confusion with (1)-(7))
add_body("【撰写要求】", indent=False, bold=True)
add_body("本部分为专利申请最重要部分，需要详细提供；发明必须是一个技术方案，不能只有原理，也不能只做功能介绍；对于软件、业务方法，要提供流程图；必须结合流程图、原理框图、电路图、系统结构图等附图进行说明。", indent=False)

add_body("结合附图，对本专利的技术方案进行描述：", indent=True, bold=True)

# (1) 总体技术路线
add_list_item("（1）总体技术路线：")
# Copy 4.1 content but skip term definitions (already in 术语解释) and the "术语定义：" heading
for p in sub41:
    text = p.text.strip()
    # Skip term definition heading
    if text == "术语定义：":
        continue
    # Skip term definition content
    if "——" in text and any(kw in text for kw in ["表面形貌状态", "脉冲历史材料状态", "材料去除状态门控", "压缩观测", "多尺度联合约束"]):
        continue
    copy_paragraph_content(p, indent=True)

# (2) 输入变量及双状态定义
add_list_item("（2）输入变量及双状态定义：")
copy_paragraphs(sub42, indent=True)

# (3) 历史状态驱动的材料去除状态门控
add_list_item("（3）历史状态驱动的材料去除状态门控：")
copy_paragraphs(sub43, indent=True)

# (4) 压缩观测与多尺度联合约束
add_list_item("（4）压缩观测与多尺度联合约束：")
copy_paragraphs(sub44, indent=True)

# Copy 压缩观测表 (Table 2 from V5)
if len(v5_tables) >= 3:
    doc.add_paragraph()
    src_table = v5_tables[2]  # 压缩观测表
    new_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    new_table.style = "Table Grid"
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell(new_table.rows[ri].cells[ci], cell.text, size=10)
    doc.add_paragraph()

# (5) 训练目标与算法实现
add_list_item("（5）训练目标与算法实现：")
copy_paragraphs(sub45, indent=True)

# Copy 模块表 (Table 3 from V5)
if len(v5_tables) >= 4:
    doc.add_paragraph()
    src_table = v5_tables[3]  # 模块表
    new_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    new_table.style = "Table Grid"
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell(new_table.rows[ri].cells[ci], cell.text, size=10)
    doc.add_paragraph()

# (6) 现有实验数据与可实施性
add_list_item("（6）现有实验数据与可实施性：")
copy_paragraphs(sub46, indent=True)

# (7) 预测输出与工程应用
add_list_item("（7）预测输出与工程应用：")
copy_paragraphs(sub47, indent=True)

# ============================================================
# 8. Section 5 - 关键点和欲保护点
# ============================================================
add_heading_text("5、本发明的关键点和欲保护点是什么？（请提炼出本发明的技术创新点，以提醒专利代理人注意，便于其撰写权利要求书）。", size=14)
copy_paragraphs(sec5, indent=True)

# ============================================================
# 9. Section 6 - 替代方案
# ============================================================
add_heading_text("6、能实现本发明目的的其他替代方案（尽量多列举可实现本发明的的其他替代方案，以便充分公开本发明）", size=14)
copy_paragraphs(sec6, indent=True)

# ============================================================
# 10. 参考文献 (template format: 【1】【2】)
# ============================================================
add_heading_text("参考文献：", size=14)
ref_num = 1
for p in references:
    text = p.text.strip()
    if text and not text.startswith("参考资料"):
        # Replace [1] with 【1】 format
        import re
        text = re.sub(r'^\[(\d+)\]', lambda m: f'【{m.group(1)}】', text)
        if not text.startswith("【"):
            text = f'【{ref_num}】{text}'
            ref_num += 1
        add_body(text, indent=False)

# ============================================================
# 11. 附录 (代理人参考材料)
# ============================================================
add_heading_text("附录：代理人参考材料", size=14)
copy_paragraphs(appendix, indent=True)

# ============================================================
# Save
# ============================================================
doc.save(V6_PATH)
print(f"V6 template-format document saved to: {V6_PATH}")
print("Done!")
