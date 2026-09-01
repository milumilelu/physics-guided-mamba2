"""Extract text from .docx and .doc files for comparison analysis."""
import re
import os
import sys

# --- Extract V4 (.docx) ---
def extract_docx(path):
    from docx import Document
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                lines.append("[TABLE] " + line)
    return "\n".join(lines)

# --- Extract Ref1 (.doc) ---
def extract_doc(path):
    import olefile
    ole = olefile.OleFileIO(path)
    # List streams
    streams = ole.listdir()
    print("Streams in .doc:", ["/".join(s) for s in streams])
    
    # Read WordDocument stream
    word_stream = ole.openstream("WordDocument").read()
    
    # Try to extract text via UTF-16LE and GBK patterns
    # Method 1: Extract all printable UTF-16LE sequences (Chinese + ASCII)
    text_parts = []
    
    # Decode entire stream as UTF-16LE with errors ignored, then clean
    try:
        decoded = word_stream.decode("utf-16-le", errors="ignore")
        # Keep CJK, ASCII printable, common punctuation, newlines
        cleaned = re.sub(r'[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\x20-\x7e\n\r\t]', ' ', decoded)
        # Collapse multiple spaces
        cleaned = re.sub(r' {3,}', '\n', cleaned)
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
        text_parts.append(cleaned)
    except Exception as e:
        print(f"UTF-16 decode error: {e}")
    
    ole.close()
    return "\n".join(text_parts)


if __name__ == "__main__":
    v4_path = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V4_工程问题导向可读性增强版.docx"
    ref_path = r"C:\Users\RZF\Desktop\博士课题资料\专利技术交底书1.doc"
    
    print("=" * 80)
    print("EXTRACTING V4 (.docx)")
    print("=" * 80)
    v4_text = extract_docx(v4_path)
    with open(r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\v4_extracted.txt", "w", encoding="utf-8") as f:
        f.write(v4_text)
    print(f"V4 extracted: {len(v4_text)} chars, {v4_text.count(chr(10))} lines")
    
    print("\n" + "=" * 80)
    print("EXTRACTING Ref1 (.doc)")
    print("=" * 80)
    ref_text = extract_doc(ref_path)
    with open(r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\ref1_extracted.txt", "w", encoding="utf-8") as f:
        f.write(ref_text)
    print(f"Ref1 extracted: {len(ref_text)} chars, {ref_text.count(chr(10))} lines")
