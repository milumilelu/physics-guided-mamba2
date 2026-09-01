"""Verify the optimized V5 document structure and content."""
from docx import Document

path = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V5_优化版.docx"
doc = Document(path)

print("=" * 70)
print("VERIFICATION REPORT")
print("=" * 70)

# 1. Count paragraphs and tables
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# 2. Check headings structure
print("\n--- DOCUMENT STRUCTURE (Headings) ---")
heading_count = 0
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith("Heading"):
        heading_count += 1
        print(f"  [{i:3d}] {para.style.name:12s} | {para.text[:80]}")
print(f"\nTotal headings: {heading_count}")

# 3. Check for ablation-related content (should be removed)
print("\n--- ABLATION CONTENT CHECK (should be empty) ---")
ablation_keywords = ["消融", "M0", "M1", "M2", "M3", "M4", "4.9", "算法验证与消融"]
found_ablation = False
for i, para in enumerate(doc.paragraphs):
    text = para.text
    for kw in ablation_keywords:
        if kw in text:
            # Exclude legitimate mentions like "M1" in other contexts
            if kw in ["M0", "M1", "M2", "M3", "M4"] and "模型" not in text and "消融" not in text:
                continue
            print(f"  WARNING [{i:3d}] keyword='{kw}': {text[:80]}")
            found_ablation = True
if not found_ablation:
    print("  PASS: No ablation-related content found.")

# Check tables for ablation
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for kw in ablation_keywords:
                if kw in cell.text:
                    print(f"  WARNING Table {ti} Row {ri} Col {ci}: keyword='{kw}': {cell.text[:60]}")
                    found_ablation = True

# 4. Check for meta-information (should be removed from main text)
print("\n--- META-INFORMATION CHECK (导读提示/撰写说明) ---")
meta_keywords = ["撰写说明", "阅读导引", "导读提示", "先看创新主线", "阅读这一节时要抓住",
                  "一句话技术路线", "理解这一节只需区分", "门控的工程意义", "为什么优先用压缩观测",
                  "两类数据各自", "训练逻辑可以理解为", "这一节的作用是说明", "本节为保护范围核心"]
found_meta = False
for i, para in enumerate(doc.paragraphs):
    text = para.text
    for kw in meta_keywords:
        if kw in text:
            print(f"  WARNING [{i:3d}] keyword='{kw}': {text[:80]}")
            found_meta = True
if not found_meta:
    print("  PASS: No meta-information found in main text.")

# 5. Check section numbering
print("\n--- SECTION NUMBERING CHECK ---")
expected_h1 = ["1、本发明要解决的技术问题", "2、技术背景及最接近现有技术",
               "3、现有技术的缺点、本发明目的及技术效果", "4、本发明技术方案的详细阐述",
               "5、本发明的关键点和欲保护点", "6、能实现本发明目的的其他替代方案",
               "附录A：代理人参考材料", "参考资料"]
h1_headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
for expected in expected_h1:
    if any(expected in h for h in h1_headings):
        print(f"  PASS: {expected}")
    else:
        print(f"  MISSING: {expected}")

# 6. Check 4.x subsections
print("\n--- 4.X SUBSECTIONS ---")
h2_headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
for h in h2_headings:
    if h.startswith("4.") or h.startswith("A.") or h.startswith("2.") or h.startswith("3."):
        print(f"  {h}")

# 7. Check term definitions in 4.1
print("\n--- TERM DEFINITIONS IN 4.1 ---")
term_keywords = ["表面形貌状态 Hₙ", "脉冲历史材料状态 Zₙ", "材料去除状态门控", "压缩观测", "多尺度联合约束"]
for kw in term_keywords:
    found = any(kw in p.text for p in doc.paragraphs)
    print(f"  {'PASS' if found else 'MISSING'}: {kw}")

# 8. Check appendix content
print("\n--- APPENDIX CONTENT ---")
appendix_keywords = ["A.2 独立权利要求技术骨架", "A.3 交代理人特别提示", "S1：", "S8：", "写作规避建议"]
for kw in appendix_keywords:
    found = any(kw in p.text for p in doc.paragraphs)
    print(f"  {'PASS' if found else 'MISSING'}: {kw}")

# 9. Check images
print("\n--- IMAGES ---")
img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_count += 1
print(f"  Total images: {img_count} (expected: 3, figure 4 removed)")

# 10. Check key point (8) modification
print("\n--- KEY POINT (8) MODIFICATION ---")
for i, para in enumerate(doc.paragraphs):
    if "（8）跨尺度验证" in para.text:
        print(f"  [{i}] {para.text}")
        if "模型消融" in para.text:
            print("  WARNING: '模型消融' still present!")
        else:
            print("  PASS: '模型消融' removed.")
        break

# 11. Word count estimate
print("\n--- CONTENT STATISTICS ---")
total_chars = sum(len(p.text) for p in doc.paragraphs)
non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
print(f"  Total characters: {total_chars}")
print(f"  Non-empty paragraphs: {non_empty}")

print("\n" + "=" * 70)
print("VERIFICATION COMPLETE")
print("=" * 70)
