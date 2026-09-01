from docx import Document

doc = Document(r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V6_模板格式版.docx")

print("=" * 80)
print("FULL CONTENT OF SECTION 4")
print("=" * 80)

in_sec4 = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "4、本发明技术方案的详细阐述" in text:
        in_sec4 = True
    if in_sec4:
        if text.startswith("5、"):
            break
        # Check for images
        from docx.oxml.ns import qn
        has_img = bool(p._element.findall('.//' + qn('w:drawing')))
        img_mark = " [IMAGE]" if has_img else ""
        bold_mark = " [BOLD]" if (p.runs and p.runs[0].bold) else ""
        if text or has_img:
            print(f"[{i:3d}]{bold_mark}{img_mark} {text[:120]}")
        else:
            print(f"[{i:3d}] (empty)")

print("\n" + "=" * 80)
print("TABLES IN SECTION 4 AREA")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    first_cell = table.rows[0].cells[0].text.strip()[:40]
    print(f"  Table {ti}: {len(table.rows)}r x {len(table.columns)}c | first: {first_cell}")
