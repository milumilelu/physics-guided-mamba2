from docx import Document

doc = Document(r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V5_优化版.docx")

print("=== 4.4 section content (end) ===")
in_44 = False
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 2" and "4.4" in p.text:
        in_44 = True
    elif p.style.name == "Heading 2" and "4.5" in p.text:
        in_44 = False
    if in_44 and p.text.strip():
        print(f"  [{i}] {p.text[:100]}")

print()
print("=== 4.5 section content (end) ===")
in_45 = False
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 2" and "4.5" in p.text:
        in_45 = True
    elif p.style.name == "Heading 2" and "4.6" in p.text:
        in_45 = False
    if in_45 and p.text.strip():
        print(f"  [{i}] {p.text[:100]}")

print()
print("=== First table (header info) ===")
if doc.tables:
    t = doc.tables[0]
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip()[:30] for c in row.cells]
        print(f"  Row {ri}: {' | '.join(cells)}")

print()
print("=== Tables count and first cell ===")
for ti, t in enumerate(doc.tables):
    first_cell = t.rows[0].cells[0].text.strip()[:40]
    print(f"  Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols | first: {first_cell}")

print()
print("=== 5 key points ===")
in_5 = False
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Heading 1" and "5、" in p.text:
        in_5 = True
    elif p.style.name == "Heading 1" and "6、" in p.text:
        in_5 = False
    if in_5 and p.text.strip() and p.style.name != "Heading 1":
        print(f"  [{i}] {p.text[:90]}")
