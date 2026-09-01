from docx import Document
from docx.oxml.ns import qn

doc = Document(r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V5_优化版.docx")

print("=== Figure references in text ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if "图1" in text or "图2" in text or "图3" in text or "图4" in text:
        print(f"  [{i}] {text[:80]}")

print()
print("=== Images in relationships ===")
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        print(f"  {rel.reltype}: {rel.target_ref}")

print()
print("=== Check inline shapes in paragraphs ===")
img_paras = 0
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall(".//" + qn("w:drawing"))
    if drawings:
        img_paras += 1
        txt = p.text[:50] if p.text else "(no text)"
        print(f"  [{i}] Paragraph with drawing: {txt}")
print(f"  Total paragraphs with drawings: {img_paras}")
