"""
Optimize V4 patent disclosure:
1. Restructure to standard 6-section format
2. Remove all ablation experiment related content
3. Clean up meta-information (导读提示, 撰写说明, etc.)
4. Renumber sections
5. Move agent reference material to appendix
"""
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V4_工程问题导向可读性增强版.docx"
DST = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V5_优化版.docx"

# Copy source to destination
shutil.copy2(SRC, DST)
doc = Document(DST)

def delete_paragraph(paragraph):
    """Delete a paragraph from the document."""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def insert_paragraph_after(paragraph, text="", style=None, bold=False):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        if bold:
            run.bold = True
    if style:
        new_para.style = style
    return new_para

def set_paragraph_text(paragraph, text, bold=None):
    """Replace all text in a paragraph, preserving style."""
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""
    # Set new text in first run, or create one
    if paragraph.runs:
        paragraph.runs[0].text = text
        if bold is not None:
            paragraph.runs[0].bold = bold
    else:
        run = paragraph.add_run(text)
        if bold is not None:
            run.bold = bold

# ============================================================
# STEP 1: Text modifications (before deletion, to use original indices)
# ============================================================
paras = doc.paragraphs

# Modify key point (8) - remove "模型消融"
# Original: （8）跨尺度验证：通过单线工况到槽工况迁移、留一频率/间距/遍次外推和模型消融证明历史状态及多尺度联合约束带来的预测改进。
p142 = paras[142]
set_paragraph_text(p142, "（8）跨尺度验证：通过单线工况到槽工况迁移、留一频率/间距/遍次外推证明历史状态及多尺度联合约束带来的预测改进。", bold=True)

# Modify section titles (renumber)
# 1、从工程问题到本发明要解决的技术问题 -> 1、本发明要解决的技术问题
set_paragraph_text(paras[20], "1、本发明要解决的技术问题")

# 4.1 总体技术路线 -> 4.1 总体技术路线与术语定义
set_paragraph_text(paras[53], "4.1 总体技术路线与术语定义")

# 4.4 单线与槽的共聚焦点云压缩观测 -> 4.4 压缩观测与多尺度联合约束
set_paragraph_text(paras[82], "4.4 压缩观测与多尺度联合约束")

# 4.6 多尺度联合反演/训练目标 -> 4.5 训练目标与算法实现
set_paragraph_text(paras[102], "4.5 训练目标与算法实现")

# 4.8 现有实验数据与可实施性 -> 4.6 现有实验数据与可实施性
set_paragraph_text(paras[116], "4.6 现有实验数据与可实施性")

# 4.10 预测输出与工程应用 -> 4.7 预测输出与工程应用
set_paragraph_text(paras[128], "4.7 预测输出与工程应用")

# 附录A title
set_paragraph_text(paras[163], "附录A：代理人参考材料")

# ============================================================
# STEP 2: Delete paragraphs (from largest index to smallest)
# ============================================================

# Indices to delete (sorted descending)
delete_indices = [
    # 交代理人特别提示 (188-193) - move to appendix
    193, 192, 191, 190, 189, 188,
    # A.2 最小算法实验 (170-176) - ablation, delete entirely
    176, 175, 174, 173, 172, 171, 170,
    # 5.1 独立权利要求骨架 (143-153) - move to appendix
    153, 152, 151, 150, 149, 148, 147, 146, 145, 144, 143,
    # 5节导读 (134)
    134,
    # 4.9 算法验证与消融实验设计 (121-127) - ablation, delete entirely
    127, 126, 125, 124, 123, 122, 121,
    # 4.8 导读 (117)
    117,
    # 4.7 标题+空行 (113-114), content (115) merges into 4.6
    114, 113,
    # 4.6 导读 (103)
    103,
    # 4.5 标题+导读+空行 (97, 98, 100), content (99, 101) merges into 4.4
    100, 98, 97,
    # 4.4 导读 (83)
    83,
    # 4.3 导读 (75)
    75,
    # 4.2 导读 (63)
    63,
    # 4.1 导读 (54)
    54,
    # 3节 导读 (39)
    39,
    # 2.2 导读 (33)
    33,
    # 0节 全部 (4-19)
    19, 18, 17, 16, 15, 14, 13, 12, 11, 10, 9, 8, 7, 6, 5, 4,
    # 撰写说明 (2)
    2,
]

# Refresh paragraph list after text modifications
paras = doc.paragraphs
for idx in delete_indices:
    if idx < len(paras):
        delete_paragraph(paras[idx])
    else:
        print(f"Warning: index {idx} out of range ({len(paras)})")

# ============================================================
# STEP 3: Delete ablation table (Table 4 - M0-M4 model comparison)
# ============================================================
tables = doc.tables
if len(tables) >= 5:
    # Table index 4 is the M0-M4 ablation table
    ablation_table = tables[4]
    tbl_element = ablation_table._element
    tbl_element.getparent().remove(tbl_element)
    print("Deleted ablation table (M0-M4)")

# ============================================================
# STEP 4: Add appendix content (A.2 independent claims + A.3 agent notes)
# Find the last paragraph of A.1 (数据整理) and insert after it
# ============================================================
paras = doc.paragraphs

# Find A.1 last item - it's the 5th bullet after "A.1 数据整理"
# After deletion, locate by text
a1_last = None
for i, p in enumerate(paras):
    if "单独标记" in p.text and "未形成稳定烧蚀" in p.text:
        a1_last = p
        break

if a1_last:
    # Insert A.2 heading
    current = insert_paragraph_after(a1_last, "A.2 独立权利要求技术骨架（供代理人参考）", style="Heading 2")
    
    # Insert description
    current = insert_paragraph_after(current, '建议优先把「模型构建/校准 + 形貌预测」写成一个完整方法，必要时再根据单一性和审查策略拆分预测方法、系统和存储介质。独权不建议出现 TRM、MoE、FNO、Koopman、spallation 等具体术语。')
    
    # Insert S1-S8
    s_items = [
        "S1：获取目标材料的超快激光加工参数及扫描/脉冲作用序列，建立初始表面形貌状态 H₀ 和历史材料状态 Z₀；",
        "S2：针对第 n 个脉冲或等效局部激光作用事件，根据当前 Hₙ、Zₙ 和激光输入 uₙ 计算至少两个材料去除状态子模型的贡献权重；",
        "S3：根据所述贡献权重组合至少两个材料去除状态子模型的输出，得到当前材料去除增量 ΔHₙ；",
        "S4：根据 ΔHₙ 更新表面形貌状态 Hₙ₊₁，并根据当前脉冲历史更新历史材料状态 Zₙ₊₁，循环直至完成预定加工过程，得到预测最终形貌；",
        "S5：针对单线加工预测形貌通过第一观测算子获得预测单线宽度和深度；针对槽加工预测形貌通过第二观测算子获得预测槽深度及至少一种粗糙度指标；",
        "S6：从单线与槽的激光共聚焦点云中按预定规则获得对应实测压缩观测，构造包含单线观测误差和槽观测误差的联合目标函数；",
        "S7：利用所述联合目标函数确定/更新历史材料状态转移参数、材料去除状态贡献参数和/或材料去除子模型参数；",
        "S8：利用确定后的模型对新的超快激光多脉冲加工条件进行形貌和/或槽深、粗糙度预测。",
    ]
    for s in s_items:
        current = insert_paragraph_after(current, s, style="List Number")
    
    # Insert writing avoidance advice
    current = insert_paragraph_after(current, '写作规避建议：独权应强调「Zₙ 与 Hₙ 为两个不同信息源」「Zₙ 由前序脉冲历史递推」「Zₙ 参与门控多个材料去除状态子模型」「两类尺度观测共同校准同一模型」。这些特征应尽量作为一个功能耦合整体出现，避免被拆解为普通逐脉冲模型 + 普通 MoE + 普通多任务学习的简单拼接。')
    
    # Insert A.3 heading
    current = insert_paragraph_after(current, "A.3 交代理人特别提示", style="Heading 2")
    
    # Insert agent notes (1-4, skip 5 which mentions ablation)
    agent_notes = [
        '1）本发明不宜将「逐脉冲形貌预测」「孵化效应」「有效脉冲数」「物理模型+残差网络」「MoE/FNO」等任一单独技术点作为核心创造性。',
        "2）独权的创造性主轴应是：独立历史材料状态 Zₙ + Zₙ 驱动去除状态门控 + 单线/槽多尺度压缩观测共同校准同一个微观模型。",
        '3）在现阶段没有时间分辨机理标签的情况下，不在独权中将门控专家写死为 spallation / phase explosion。使用「材料去除状态子模型」更稳妥。',
        '4）实验数据尚未证明某一微观状态的唯一物理含义，因此说明书应使用「联合约束、反演标定、一致性验证」等表述，避免使用「由宏观粗糙度证明某微观机理」等过强因果表述。',
    ]
    for note in agent_notes:
        current = insert_paragraph_after(current, note)
    
    print("Added appendix A.2 and A.3")
else:
    print("Warning: Could not find A.1 last paragraph")

# ============================================================
# STEP 5: Add term definitions at the beginning of 4.1
# Find "第一步" paragraph in 4.1 and insert before it
# ============================================================
paras = doc.paragraphs
first_step = None
for i, p in enumerate(paras):
    if "第一步" in p.text and "描述" in p.text:
        first_step = p
        break

if first_step:
    # Insert term definitions before 第一步
    # We need to insert before, so find the paragraph before first_step
    # Actually, insert after the paragraph before first_step
    # Simpler: insert after first_step's previous sibling
    
    # Get the paragraph before first_step
    prev_para = None
    for i, p in enumerate(paras):
        if p._element is first_step._element:
            if i > 0:
                prev_para = paras[i-1]
            break
    
    if prev_para:
        # Insert term definition heading
        current = insert_paragraph_after(prev_para, "术语定义：", bold=True)
        current = insert_paragraph_after(current, "表面形貌状态 Hₙ——第 n 个脉冲作用前的材料表面几何状态，可采用局部高度场、深度轮廓、坡度/曲率或其降维表征。")
        current = insert_paragraph_after(current, "脉冲历史材料状态 Zₙ——由先前脉冲序列造成、且不能由当前表面形貌 Hₙ 唯一确定的材料历史状态，可表征材料弱化、缺陷累积、残余热状态、光学性质变化或其隐变量组合。")
        current = insert_paragraph_after(current, "材料去除状态门控——依据 Hₙ、Zₙ 与当前脉冲条件计算多个材料去除状态子模型的贡献权重，实现不同加工阶段去除响应的软切换或加权组合。")
        current = insert_paragraph_after(current, "压缩观测——由激光共聚焦点云按统一几何/统计规则得到的低维实验量，单线观测包括线宽、线深，槽观测包括平均/特征槽深以及 Sa、Sq、Sz 等粗糙度指标。")
        current = insert_paragraph_after(current, "多尺度联合约束——使用同一微观状态模型同时生成单线与槽加工结果，并通过两类不同尺度的压缩实验观测共同确定历史状态演化参数和机制门控参数。")
        print("Added term definitions to 4.1")
    else:
        print("Warning: Could not find previous paragraph for term definitions")
else:
    print("Warning: Could not find 第一步 paragraph")

# ============================================================
# Save
# ============================================================
doc.save(DST)
print(f"\nSaved optimized document to: {DST}")
print("Done!")
