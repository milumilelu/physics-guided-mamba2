"""
Convert V5 optimized patent disclosure to V7 template format.
Changes from V6:
1. Remove section 2.3 (中国人工智能相关专利撰写考虑)
2. Reorder section 4: (1) 专利数据需求定义 (moved from old (6)), then (2)-(7)
3. Insert 6 new patent figures (fig1-fig6 PNG) at appropriate positions
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.parts.image import ImagePart
import copy
import os

V5_PATH = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V5_优化版.docx"
V7_PATH = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V7_最终版.docx"
FIG_DIR = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\figures_png"

# Read V5 content
v5 = Document(V5_PATH)
v5_paras = v5.paragraphs
v5_tables = v5.tables

def extract_section(paras, start_heading, end_heading=None):
    result = []
    in_section = False
    for p in paras:
        if p.style.name == "Heading 1" and start_heading in p.text:
            in_section = True
            continue
        if in_section:
            if end_heading and p.style.name == "Heading 1" and end_heading in p.text:
                break
            if p.style.name == "Heading 1" and p.text != start_heading:
                break
            result.append(p)
    return result

def extract_subsection(paras, heading_text):
    result = []
    in_section = False
    for p in paras:
        if p.style.name == "Heading 2" and heading_text in p.text:
            in_section = True
            continue
        if in_section:
            if p.style.name in ("Heading 1", "Heading 2"):
                break
            result.append(p)
    return result

# Extract V5 sections
sec1 = extract_section(v5_paras, "1、本发明要解决的技术问题")
sec2 = extract_section(v5_paras, "2、技术背景及最接近现有技术")
sec3 = extract_section(v5_paras, "3、现有技术的缺点")
sec5 = extract_section(v5_paras, "5、本发明的关键点和欲保护点")
sec6 = extract_section(v5_paras, "6、能实现本发明目的的其他替代方案")
appendix = extract_section(v5_paras, "附录A：代理人参考材料")
references = extract_section(v5_paras, "参考资料")

# Extract 4.x subsections
sub41 = extract_subsection(v5_paras, "4.1 总体技术路线与术语定义")
sub42 = extract_subsection(v5_paras, "4.2 输入变量及双状态定义")
sub43 = extract_subsection(v5_paras, "4.3 历史状态驱动的材料去除状态门控")
sub44 = extract_subsection(v5_paras, "4.4 压缩观测与多尺度联合约束")
sub45 = extract_subsection(v5_paras, "4.5 训练目标与算法实现")
sub46 = extract_subsection(v5_paras, "4.6 现有实验数据与可实施性")
sub47 = extract_subsection(v5_paras, "4.7 预测输出与工程应用")

# Extract 2.x subsections (excluding 2.3)
sub21 = extract_subsection(v5_paras, "2.1 超快激光多脉冲烧蚀形貌预测的工程需求")
sub22 = extract_subsection(v5_paras, "2.2 与本发明最接近的现有技术方案")

# Extract term definitions from 4.1
term_defs = []
for p in sub41:
    if "——" in p.text and any(kw in p.text for kw in ["表面形貌状态", "脉冲历史材料状态", "材料去除状态门控", "压缩观测", "多尺度联合约束"]):
        term_defs.append(p.text)

# ============================================================
# Create V7 document
# ============================================================
doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "宋体"
font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_title(text, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading_text(text, size=14, space_before=12, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_body(text, indent=True, bold=False, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_list_item(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_figure(fig_name, caption_text, width_cm=14):
    """Insert a figure with caption."""
    fig_path = os.path.join(FIG_DIR, fig_name)
    if not os.path.exists(fig_path):
        print(f"WARNING: Figure not found: {fig_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(fig_path, width=Cm(width_cm))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run(caption_text)
    cap_run.bold = True
    cap_run.font.size = Pt(10.5)
    cap_run.font.name = "宋体"
    cap_run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

def copy_paragraph_content(src_para, indent=True):
    has_content = bool(src_para.text.strip()) or bool(src_para._element.findall('.//' + qn('w:drawing')))
    if not has_content:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        return p
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    for child in src_para._element:
        if child.tag == qn('w:pPr'):
            continue
        new_child = copy.deepcopy(child)
        p._element.append(new_child)
    for blip in p._element.findall('.//' + qn('a:blip')):
        old_rId = blip.get(qn('r:embed'))
        if old_rId and old_rId in src_para.part.rels:
            src_rel = src_para.part.rels[old_rId]
            src_image_part = src_rel.target_part
            new_image_part = ImagePart(
                src_image_part.partname,
                src_image_part.content_type,
                src_image_part.blob,
                doc.part.package
            )
            new_rId = doc.part.relate_to(new_image_part, src_rel.reltype)
            blip.set(qn('r:embed'), new_rId)
    for run in p.runs:
        if run.font.name is None:
            run.font.name = "宋体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if run.font.size is None:
            run.font.size = Pt(12)
    return p

def copy_paragraphs(paras, indent=True, skip_images=False):
    for p in paras:
        # Skip image-only paragraphs and their captions when skip_images=True
        if skip_images:
            has_drawing = bool(p._element.findall('.//' + qn('w:drawing')))
            text = p.text.strip()
            # Skip image-only paragraphs
            if has_drawing and not text:
                continue
            # Skip figure captions (old V5 figures)
            if text.startswith("图") and any(kw in text for kw in ["微观双状态", "双状态—机制门控", "共聚焦点云压缩"]):
                continue
        copy_paragraph_content(p, indent=indent)

def set_cell(cell, text, size=10, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.bold = bold

# ============================================================
# 1. Title
# ============================================================
add_title("专 利 申 请 技 术 交 底 书", size=22, space_after=16)

# ============================================================
# 2. Header info table
# ============================================================
info_table = doc.add_table(rows=10, cols=4)
info_table.style = "Table Grid"
info_data = [
    ["发明类型：", "发明 ☑  实用新型 ☐  外观设计 ☐", "", ""],
    ["同日申请实用新型：", "☐", "请求提前公布：", "☐"],
    ["优先权 实审方式：", "☐ 新申请提  ☐ 公布提  ☐ 期限届满提  ☐ 不提", "", ""],
    ["发明名称：", "一种基于双状态机制门控及多尺度观测约束的超快激光多脉冲烧蚀形貌预测方法", "", ""],
    ["发明人：", "（待填写）", "身份证号：", "（待填写）"],
    ["技术联系人：", "（待填写）", "所属公司/部门：", "（待填写）"],
    ["电话：", "（待填写）", "传真：", "（待填写）"],
    ["E-mail：", "（待填写）", "", ""],
    ["版本说明：", "V7：最终版（删除2.3节，数据需求前置，插入6张专利附图）", "", ""],
    ["", "", "", ""],
]
for ri, row_data in enumerate(info_data):
    for ci, cell_text in enumerate(row_data):
        set_cell(info_table.rows[ri].cells[ci], cell_text, size=10)

doc.add_paragraph()

# ============================================================
# 3. 术语解释
# ============================================================
add_heading_text("术语解释：", size=14, space_before=8)
for td in term_defs:
    add_body(td, indent=True)

# ============================================================
# 4. Section 1 - 技术问题
# ============================================================
add_heading_text("1、本发明要解决的技术问题是什么？", size=14)
copy_paragraphs(sec1, indent=True)
# Insert fig1 at end of section 1
add_figure("fig1.png", "图1  多脉冲扫描加工过程与宏观形貌形成")

# ============================================================
# 5. Section 2 - 技术背景 (only 2.1 and 2.2, NO 2.3)
# ============================================================
add_heading_text("2、详细介绍技术背景,并描述已有的与本发明最相近似的实现方案（包括两部分：1、作为本发明基础的且帮助理解本发明公知技术内容；2、与本发明最接近的已有技术方案的说明）。", size=14)

add_heading_text("2.1 超快激光多脉冲烧蚀形貌预测的工程需求", size=12, space_before=8)
copy_paragraphs(sub21, indent=True)

add_heading_text("2.2 与本发明最接近的现有技术方案", size=12, space_before=8)
copy_paragraphs(sub22, indent=True)

# Copy 查新对比表 (Table 1 from V5)
if len(v5_tables) >= 2:
    doc.add_paragraph()
    src_table = v5_tables[1]
    new_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    new_table.style = "Table Grid"
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell(new_table.rows[ri].cells[ci], cell.text, size=10)
    doc.add_paragraph()

# NOTE: 2.3 中国人工智能相关专利撰写考虑 is DELETED per user request

# ============================================================
# 6. Section 3 - 缺点/目的/效果
# ============================================================
add_heading_text("3、以因果关系推理的方式推导出现有技术的缺点是什么？针对这些缺点，说明本发明的目的及能够达到的技术效果。（现有技术的缺点是针对于本发明的优点来说的，本发明不能解决的缺点不必写）", size=14)
copy_paragraphs(sec3, indent=True, skip_images=True)
# Insert fig2 and fig3 at end of section 3
add_figure("fig2.png", "图2  同一表面形貌可由不同历史材料状态达到")
add_figure("fig3.png", "图3  单线与槽观测对脉冲级模型的互补约束")

# ============================================================
# 7. Section 4 - 技术方案详细阐述 (reordered)
# ============================================================
add_heading_text("4、本发明技术方案的详细阐述，应该结合流程图、原理框图、电路图、系统结构图进行说明（发明中每一功能的实现都要有相应的技术实现方案；所有附图都应该有详细的文字描述；方法专利都应该提供流程图，并提供相关的系统装置）。", size=14)

add_body("【撰写要求】", indent=False, bold=True)
add_body("本部分为专利申请最重要部分，需要详细提供；发明必须是一个技术方案，不能只有原理，也不能只做功能介绍；对于软件、业务方法，要提供流程图；必须结合流程图、原理框图、电路图、系统结构图等附图进行说明。", indent=False)

add_body("结合附图，对本专利的技术方案进行描述：", indent=True, bold=True)

# (1) 专利数据需求定义 (moved from old (6), reorganized)
add_list_item("（1）专利数据需求定义：")
add_body("本发明的训练与验证数据需满足以下定义，以确保模型参数可被有效反演校准，并保持工程应用的低数据成本。", indent=True)
# Copy content from old sub46 (现有实验数据与可实施性)
for p in sub46:
    copy_paragraph_content(p, indent=True)

# (2) 总体技术路线 (old (1))
add_list_item("（2）总体技术路线：")
for p in sub41:
    text = p.text.strip()
    if text == "术语定义：":
        continue
    if "——" in text and any(kw in text for kw in ["表面形貌状态", "脉冲历史材料状态", "材料去除状态门控", "压缩观测", "多尺度联合约束"]):
        continue
    # Skip old V5 image and its caption
    has_drawing = bool(p._element.findall('.//' + qn('w:drawing')))
    if has_drawing and not text:
        continue
    if text.startswith("图") and "双状态—机制门控" in text:
        continue
    copy_paragraph_content(p, indent=True)
# Insert fig4 at end of (2)
add_figure("fig4.png", "图4  基于双状态递推与多尺度观测联合约束的总体技术路线")

# (3) 输入变量及双状态定义 (old (2))
add_list_item("（3）输入变量及双状态定义：")
copy_paragraphs(sub42, indent=True)

# (4) 历史状态驱动的材料去除状态门控 (old (3))
add_list_item("（4）历史状态驱动的材料去除状态门控：")
copy_paragraphs(sub43, indent=True)
# Insert fig5 at end of (4)
add_figure("fig5.png", "图5  单脉冲内部计算流水线")

# (5) 压缩观测与多尺度联合约束 (old (4))
add_list_item("（5）压缩观测与多尺度联合约束：")
copy_paragraphs(sub44, indent=True, skip_images=True)
# Copy 压缩观测表 (Table 2 from V5)
if len(v5_tables) >= 3:
    doc.add_paragraph()
    src_table = v5_tables[2]
    new_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    new_table.style = "Table Grid"
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell(new_table.rows[ri].cells[ci], cell.text, size=10)
    doc.add_paragraph()

# (6) 训练目标与算法实现 (old (5))
add_list_item("（6）训练目标与算法实现：")
copy_paragraphs(sub45, indent=True)
# Insert fig6 at end of (6)
add_figure("fig6.png", "图6  单线与槽多尺度压缩观测联合反演闭环")
# Copy 模块表 (Table 3 from V5)
if len(v5_tables) >= 4:
    doc.add_paragraph()
    src_table = v5_tables[3]
    new_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    new_table.style = "Table Grid"
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell(new_table.rows[ri].cells[ci], cell.text, size=10)
    doc.add_paragraph()

# (7) 预测输出与工程应用 (old (7))
add_list_item("（7）预测输出与工程应用：")
copy_paragraphs(sub47, indent=True)

# ============================================================
# 8. Section 5 - 关键点
# ============================================================
add_heading_text("5、本发明的关键点和欲保护点是什么？（请提炼出本发明的技术创新点，以提醒专利代理人注意，便于其撰写权利要求书）。", size=14)
copy_paragraphs(sec5, indent=True)

# ============================================================
# 9. Section 6 - 替代方案
# ============================================================
add_heading_text("6、能实现本发明目的的其他替代方案（尽量多列举可实现本发明的的其他替代方案，以便充分公开本发明）", size=14)
copy_paragraphs(sec6, indent=True)

# ============================================================
# 10. 参考文献
# ============================================================
add_heading_text("参考文献：", size=14)
ref_num = 1
for p in references:
    text = p.text.strip()
    if text and not text.startswith("参考资料"):
        import re
        text = re.sub(r'^\[(\d+)\]', lambda m: f'【{m.group(1)}】', text)
        if not text.startswith("【"):
            text = f'【{ref_num}】{text}'
            ref_num += 1
        add_body(text, indent=False)

# ============================================================
# 11. 附录
# ============================================================
add_heading_text("附录：代理人参考材料", size=14)
copy_paragraphs(appendix, indent=True)

# ============================================================
# Save
# ============================================================
doc.save(V7_PATH)
print(f"V7 final document saved to: {V7_PATH}")
print("Done!")
