"""Read full structure of V4 docx: paragraphs with styles, tables, images."""
from docx import Document
from docx.oxml.ns import qn

path = r"C:\Users\RZF\Desktop\博士课题资料\physics-guided Mamba-2\专利\超快激光多脉冲烧蚀形貌预测_专利技术交底书_V4_工程问题导向可读性增强版.docx"
doc = Document(path)

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else "None"
    text = para.text.strip()
    if text:
        # Check for runs with bold/italic
        has_bold = any(run.bold for run in para.runs if run.bold)
        print(f"[{i:3d}] style={style:20s} bold={has_bold} | {text[:120]}")
    else:
        print(f"[{i:3d}] style={style:20s} | (empty)")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:40] for cell in row.cells]
        print(f"  Row {ri}: {' | '.join(cells)}")

print("\n=== SECTIONS ===")
for si, section in enumerate(doc.sections):
    print(f"Section {si}: page_width={section.page_width}, page_height={section.page_height}")

# Count images
print("\n=== IMAGES ===")
img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_count += 1
        print(f"  Image: {rel.target_ref}")
print(f"Total images: {img_count}")
